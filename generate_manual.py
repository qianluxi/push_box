# -*- coding: utf-8 -*-
"""Generate push_box_game_manual.docx using python-docx."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn('w:shd'), {
            qn('w:fill'): hex_color,
            qn('w:val'): 'clear',
        }
    )
    shading_elm.append(shading)


def add_cell_text(cell, text, bold=False, size=12, font_name='Microsoft YaHei', color=(0, 0, 0)):
    """Add text to a table cell with formatting."""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_ele = run._element
    r_ele.set(qn('w:eastAsiaFont'), font_name)
    run.font.color.rgb = RGBColor(*color)


doc = Document()

# Set default fonts
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# Colors
ACCENT_COLOR = RGBColor(0xD9, 0x77, 0x06)
SURFACE_BG = "FEF3C7"
PRIMARY_COLOR = RGBColor(0x0F, 0x17, 0x2A)
BODY_COLOR = RGBColor(0x1E, 0x29, 0x3B)
SECONDARY_COLOR = RGBColor(0x64, 0x74, 0x8B)


def add_heading_custom(doc, text, level=1):
    """Add heading with custom styling."""
    if level == 1:
        p = doc.add_heading(text, level=1)
    else:
        p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = PRIMARY_COLOR
        run.font.name = 'Microsoft YaHei'
        run.bold = True
        r_e = run._element
        r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')
    return p


def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.71)  # 2 Chinese chars
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Microsoft YaHei'
    r_e = run._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')
    return p


def add_note_box(doc, label, content):
    """Add a note box as a bordered table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set border on the single cell
    cell = t.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    for border_name in ['top', 'bottom', 'left', 'right']:
        border_elm = tc_pr.makeelement(
            qn(f'w:{border_name}'),
            {qn('w:val'): 'single', qn('w:sz'): '8', qn('w:space'): '0', qn('w:color'): 'D97706'}
        )
        tc_pr.append(border_elm)
    # Shading on the cell
    set_cell_shading(cell, SURFACE_BG)
    cell.paragraphs[0].clear()
    run_label = cell.paragraphs[0].add_run(label + " ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.name = 'Microsoft YaHei'
    run_label.font.color.rgb = ACCENT_COLOR
    r_e = run_label._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')
    run_content = cell.paragraphs[0].add_run(content)
    run_content.font.size = Pt(10)
    run_content.font.name = 'Microsoft YaHei'
    run_content.font.color.rgb = BODY_COLOR
    r_e2 = run_content._element
    r_e2.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')


def add_key_combo(doc, keys_str):
    """Add styled key combo."""
    p = doc.add_paragraph()
    keys = keys_str.split(", ")
    for i, key in enumerate(keys):
        run = p.add_run("[" + key + "]")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = ACCENT_COLOR
        r_e = run._element
        r_e.set(qn('w:eastAsiaFont'), 'Calibri')
        if i < len(keys) - 1:
            p.add_run("  ")


# ================================================================
# COVER PAGE
# ================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)  # Approximate top spacing

p_run = p.add_run("推箱子游戏手册")
p_run.font.size = Pt(28)
p_run.bold = True
p_run.font.name = 'Microsoft YaHei'
p_run.font.color.rgb = PRIMARY_COLOR
p_run._element.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2_run = p2.add_run("Sokoban Game Manual -- Version 1.1")
p2_run.font.size = Pt(12)
p2_run.font.name = 'Calibri'
p2_run.font.color.rgb = SECONDARY_COLOR
p2_run._element.set(qn('w:eastAsiaFont'), 'Calibri')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3_run = p3.add_run("\u2014\u2014 V1.1 \u7248\u672c\u73a9\u5bb6\u6307\u5357")
p3_run.font.size = Pt(11)
p3_run.font.name = 'Microsoft YaHei'
p3_run.font.color.rgb = SECONDARY_COLOR
p3_run._element.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

# Spacer
doc.add_paragraph()

# Meta info
meta_items = [
    ("版本", "V1.1"),
    ("适配", "Pygame 2.6.1+ / Python 3.12+"),
    ("作者", "qianluxi"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run_label = p.add_run(label + ": ")
    run_label.font.size = Pt(10)
    run_label.font.name = 'Microsoft YaHei'
    run_label.font.color.rgb = SECONDARY_COLOR
    r_e = run_label._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)
    run_val.font.name = 'Microsoft YaHei'
    run_val.font.color.rgb = SECONDARY_COLOR
    r_e2 = run_val._element
    r_e2.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

# Bottom spacer
doc.add_paragraph()

# Footer line on cover
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f1 = p_footer.add_run("Sokoban 推箱子游戏")
run_f1.font.size = Pt(9)
run_f1.font.name = 'Microsoft YaHei'
run_f1.font.color.rgb = SECONDARY_COLOR
p_footer.add_run("          ")
run_f2 = p_footer.add_run("github.com/qianluxi/push_box")
run_f2.font.size = Pt(9)
run_f2.font.name = 'Calibri'
run_f2.font.color.rgb = SECONDARY_COLOR

# Page break after cover
doc.add_page_break()

# ================================================================
# CONTENT PAGES
# ================================================================

# Inner title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p.add_run("推箱子游戏手册")
run_t.font.size = Pt(22)
run_t.bold = True
run_t.font.name = 'Microsoft YaHei'
run_t.font.color.rgb = PRIMARY_COLOR

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t2 = p2.add_run("Sokoban Game Manual -- Version 1.1")
run_t2.font.size = Pt(11)
run_t2.font.name = 'Calibri'
run_t2.font.color.rgb = SECONDARY_COLOR

# Divider
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_div = p_div.add_run("------")
run_div.font.size = Pt(10)
run_div.font.color.rgb = ACCENT_COLOR
p_div.paragraph_format.space_before = Pt(12)
p_div.paragraph_format.space_after = Pt(12)

# Welcome message
p_w = doc.add_paragraph()
p_w.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_w = p_w.add_run("欢迎阅读本手册！这里包含了游戏的全部操作说明、规则详解和进阶技巧。无论你是初识推箱子的新手，还是想提升水平的老玩家，都能在这里找到有用的信息。")
run_w.font.size = Pt(11)
run_w.font.name = 'Microsoft YaHei'

doc.add_paragraph()

# Chapter 1
add_heading_custom(doc, "第一章  游戏简介", 1)
p_div2 = doc.add_paragraph()
p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = p_div2.add_run("------")
run_d.font.color.rgb = ACCENT_COLOR
p_div2.paragraph_format.space_before = Pt(6)
p_div2.paragraph_format.space_after = Pt(6)

add_body(doc, "推箱子（Sokoban）是一种经典的智力推理游戏，最初由Johan Stevanne于1981年创作。玩家需要将所有箱子推到指定目标位置即可关当前关卡。游戏看似简单，但难度逐渐上升，需要较好的空间思维和策略规划能力。")

add_body(doc, "本项目是基于Python + Pygame构建的现代版推箱子移动游戏，V1.1版本完全重构了核心代码架构，提升了稳定性和可维护性。")

add_note_box(doc, "核心特点", "多关卡设计，难度逐渐增加，包含经典手动关卡和自助编辑关卡")
add_note_box(doc, "技术特色", "干净的的分层架构（Board/State/Rules/渲染），支持Undo撤回和关卡切换")
add_note_box(doc, "系统要求", "Python 3.12+，Pygame 2.6.1+，支持Windows / macOS / Linux")

doc.add_paragraph()

# Chapter 2
add_heading_custom(doc, "第二章  操作指南", 1)
p_div2 = doc.add_paragraph()
p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d2 = p_div2.add_run("------")
run_d2.font.color.rgb = ACCENT_COLOR
p_div2.paragraph_format.space_before = Pt(6)
p_div2.paragraph_format.space_after = Pt(6)

add_heading_custom(doc, "一  移动控制", 2)
add_body(doc, "使用方向键或WASD键控制玩家移动：")

# Control table
t = doc.add_table(rows=4, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = True

# Header row
for i, text in enumerate(["按键", "功能", "说明"]):
    cell = t.rows[0].cells[i]
    set_cell_shading(cell, SURFACE_BG)
    for para in cell.paragraphs:
        para.clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = PRIMARY_COLOR
    r_e = run._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

# Data rows
row_data = [
    (["\u2191 \u2193 \u2190 \u2192", ""], "移动", "上下左右方向移动玩家"),
    (["W A S D", ""], "移动", "方向键的替代方式"),
]
keys_row_data = [
    (["上\n下\n左\n右"], "移动", "上下左右方向移动玩家"),
    (["W\nA\nS\nD"], "移动", "方向键的替代方式"),
]

for row_idx, keys in enumerate(keys_row_data, start=1):
    for col_idx, val in enumerate(keys):
        cell = t.rows[row_idx].cells[col_idx]
        for para in cell.paragraphs:
            para.clear()
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = BODY_COLOR
        r_e = run._element
        r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

add_heading_custom(doc, "二  功能按键", 2)

# Function keys table
t2 = doc.add_table(rows=7, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(["按键", "功能"]):
    cell = t2.rows[0].cells[i]
    set_cell_shading(cell, SURFACE_BG)
    for para in cell.paragraphs:
        para.clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = PRIMARY_COLOR
    r_e = run._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

func_data = [
    ("Z, Ctrl+Z", "撤回一步"),
    ("R", "重置关卡"),
    ("ESC", "暂停游戏"),
    ("Enter, Space", "确认选择 / 进入下一关"),
    ("Tab", "上一关 / 下一关"),
    ("Q", "退出游戏"),
]
for row_idx, (key, func) in enumerate(func_data, start=1):
    cell_k = t2.rows[row_idx].cells[0]
    cell_f = t2.rows[row_idx].cells[1]
    for para in cell_k.paragraphs:
        para.clear()
    run_k = cell_k.paragraphs[0].add_run(key)
    run_k.bold = True
    run_k.font.size = Pt(10)
    run_k.font.name = 'Calibri'
    run_k.font.color.rgb = ACCENT_COLOR
    r_k = run_k._element
    r_k.set(qn('w:eastAsiaFont'), 'Calibri')

    for para in cell_f.paragraphs:
        para.clear()
    run_f = cell_f.paragraphs[0].add_run(func)
    run_f.font.size = Pt(10)
    run_f.font.name = 'Microsoft YaHei'
    run_f.font.color.rgb = BODY_COLOR
    r_f = run_f._element
    r_f.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

add_heading_custom(doc, "三  暂停菜单", 2)
add_body(doc, "按 ESC 进入暂停状态，在此状态下可以使用以下操作：")
add_note_box(doc, "重置关卡", "按 R 回到当前关卡初始状态")
add_note_box(doc, "继续游戏", "按 ESC 或 Enter 退出暂停状态")
add_note_box(doc, "退出游戏", "按 Q 退出回主菜单或结束游戏")

doc.add_paragraph()

# Chapter 3
add_heading_custom(doc, "第三章  游戏规则", 1)
p_div3 = doc.add_paragraph()
p_div3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d3 = p_div3.add_run("------")
run_d3.font.color.rgb = ACCENT_COLOR
p_div3.paragraph_format.space_before = Pt(6)
p_div3.paragraph_format.space_after = Pt(6)

add_heading_custom(doc, "一  目标", 2)
add_body(doc, "在每一关中，你的任务是将所有箱子推到目标位置。当所有箱子都在目标位置上时，关卡即为通关。")

add_heading_custom(doc, "二  地图元素", 2)
add_body(doc, "关卡地图由以下元素组成：")

# Elements table
t3 = doc.add_table(rows=8, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
symbols = ["#", "空格", ".", "$", "@", "*", "+"]
meanings = ["墙", "地板", "目标", "箱子", "玩家", "箱子在目标上", "玩家在目标上"]
explanations = [
    "不可通行的固定障碍物",
    "可行走的空地",
    "箱子需要推到的位置",
    "需要推移的可移动物体",
    "你控制的角色",
    "已经达标的箱子",
    "玩家站在目标位置上",
]
for i, text in enumerate(["符号", "含义", "说明"]):
    cell = t3.rows[0].cells[i]
    set_cell_shading(cell, SURFACE_BG)
    for para in cell.paragraphs:
        para.clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = PRIMARY_COLOR
    r_e = run._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

# Accent colors for special symbols
for row_idx in range(1, 8):
    sym = symbols[row_idx - 1]
    meaning = meanings[row_idx - 1]
    explanation = explanations[row_idx - 1]

    for col_idx, val in enumerate([sym, meaning, explanation]):
        cell = t3.rows[row_idx].cells[col_idx]
        for para in cell.paragraphs:
            para.clear()
        run = cell.paragraphs[0].add_run(val)
        if col_idx == 0 and sym in ("*", "+"):
            run.font.color.rgb = ACCENT_COLOR
        elif col_idx == 0:
            run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = BODY_COLOR
        r_e = run._element
        r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

add_heading_custom(doc, "三  移动规则", 2)
add_body(doc, "一  玩家可以自由走到任意无墙的地板方向；")
add_body(doc, "二  玩家可以在箱子后方，将箱子推到前方的地板上（包括在目标位置上）；")
add_body(doc, "三  箱子只能被推，不能被拉；")
add_body(doc, "四  玩家无法穿越墙、箱子或其他障碍物；")
add_body(doc, "五  如果箱子前方是墙或另一个箱子，则无法推动，提示 BLOCKED（被阻止）。")

add_heading_custom(doc, "四  关卡解关条件", 2)
add_body(doc, "当且仅当所有箱子都处于目标位置时，关卡才视为完成。如果某个箱子已经在目标上，你可以将它推开，但最终所有箱子都必须回到目标位置。")
add_note_box(doc, "警告", "一旦将箱子推到角落或两面墙之间的死角位置，可能无法再移动。不确定的推法请先按 Z 撤回。")

doc.add_paragraph()

# Chapter 4
add_heading_custom(doc, "第四章  关卡说明", 1)
p_div4 = doc.add_paragraph()
p_div4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d4 = p_div4.add_run("------")
run_d4.font.color.rgb = ACCENT_COLOR
p_div4.paragraph_format.space_before = Pt(6)
p_div4.paragraph_format.space_after = Pt(6)

add_body(doc, "本游戏预置了 7 关经典推箱子关卡，难度从简单到困难逐渐上升。你可以使用 Tab 键在关卡之间切换。")

add_heading_custom(doc, "关卡列表", 2)

# Level table
t4 = doc.add_table(rows=8, cols=4)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(["关卡", "名称", "难度", "特点"]):
    cell = t4.rows[0].cells[i]
    set_cell_shading(cell, SURFACE_BG)
    for para in cell.paragraphs:
        para.clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = PRIMARY_COLOR
    r_e = run._element
    r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

levels = [
    ("Level 1", "初识推箱子", "\u2605\u2606\u2606\u2606\u2606", "单个箱子，位置明显"),
    ("Level 2", "转弯思维", "\u2605\u2605\u2606\u2606\u2606", "需要转弯才能推到目标"),
    ("Level 3", "多箱子挑战", "\u2605\u2605\u2605\u2606\u2606", "多个箱子需要协调"),
    ("Level 4", "独立思考", "\u2605\u2605\u2605\u2605\u2606", "有障碍的干扰关卡"),
    ("Level 5", "策略推理", "\u2605\u2605\u2605\u2605\u2606", "要求一定的推理顺序"),
    ("Level 6", "复杂布局", "\u2605\u2605\u2605\u2605\u2605", "复杂走廊，需要多步预案"),
    ("Level 7", "经典挑战", "\u2605\u2605\u2605\u2605\u2605", "难度最高，需要全面运用所学技巧"),
]
for row_idx, (level, name, diff, feat) in enumerate(levels, start=1):
    vals = [level, name, diff, feat]
    for col_idx, val in enumerate(vals):
        cell = t4.rows[row_idx].cells[col_idx]
        for para in cell.paragraphs:
            para.clear()
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        if col_idx == 2:
            run.font.color.rgb = ACCENT_COLOR
        else:
            run.font.color.rgb = BODY_COLOR
        r_e = run._element
        r_e.set(qn('w:eastAsiaFont'), 'Microsoft YaHei')

add_heading_custom(doc, "创建自定义关卡", 2)
add_body(doc, "你可以在 levels/ 目录下创建自己的关卡文件（txt格式），并按照命名规则将其加入游戏：")

# Example code block
t_example = doc.add_table(rows=1, cols=1)
t_example.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t_example.cell(0, 0)
set_cell_shading(cell, SURFACE_BG)
for para in cell.paragraphs:
    para.clear()
code_lines = "# 关卡文件格式示例:", "########", "#      #", "#  .   #", "#  $   #", "#  @   #", "########"
run_code = cell.paragraphs[0].add_run("\n".join(code_lines))
run_code.font.size = Pt(10)
run_code.font.name = 'Courier New'
run_code.font.color.rgb = BODY_COLOR

add_body(doc, "在文件中上述字符构建你的关卡布局。文件名建议使用 level_XX.txt 格式，其中 XX 为数字（如 level_08.txt），用于确定关卡顺序。")

doc.add_paragraph()

# Chapter 5
add_heading_custom(doc, "第五章  技巧与策略", 1)
p_div5 = doc.add_paragraph()
p_div5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d5 = p_div5.add_run("------")
run_d5.font.color.rgb = ACCENT_COLOR
p_div5.paragraph_format.space_before = Pt(6)
p_div5.paragraph_format.space_after = Pt(6)

add_heading_custom(doc, "一  新手建议", 2)
add_body(doc, "一  从第一关开始，不要急于推箱子。先观察全局布局，想好策略再行动。")
add_body(doc, "二  养成使用 Undo 的习惯。不确定的操作一定要按 Z 撤回，不要怕花时间。")
add_body(doc, "三  记住关卡的目标位置和箱子数量，这帮助了你判断进展。")

add_heading_custom(doc, "二  中级策略", 2)
add_note_box(doc, "避免角落陷阱", "将箱子推进两面墙交叉的角落是最常见的失败原因。一旦箱子贴墙，确保少数情况下才推。")
add_note_box(doc, "一次只推一个箱子", "如果两个箱子靠得很近，将它们推在一起会极大限制你的操作空间。")
add_note_box(doc, "先推远的，再推近的", "当目标位置较远时，先处理远处的箱子可以避免它们阻挡路径。")

add_heading_custom(doc, "三  高级技巧", 2)
add_note_box(doc, "回旋法", "当箱子推错了，需要将它推回原位重新推。保持足够的操作空间是关键。")
add_note_box(doc, "路径预览", "在推箱子之前，先移动到玩家的正确位置。有时你需要绕一大圈才能得到最佳角度。")
add_note_box(doc, "利用目标位作为假路径", "有时目标位本身是一条路径上的通道，你可以先把箱子推过目标再推回。")
add_note_box(doc, "经验总结", "高手通常会在按下方向键之前，在心里模拟整个推箱流程。如果一步错了就完不成，就不要推。")

doc.add_paragraph()

# Chapter 6
add_heading_custom(doc, "第六章  常见问题", 1)
p_div6 = doc.add_paragraph()
p_div6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d6 = p_div6.add_run("------")
run_d6.font.color.rgb = ACCENT_COLOR
p_div6.paragraph_format.space_before = Pt(6)
p_div6.paragraph_format.space_after = Pt(6)

add_note_box(doc, "游戏打不开", "确保已安装 Python 3.12+ 和 Pygame。在游戏目录下运行\"pip install pygame\"安装依赖，然后执行\"python main.py\"。")
add_note_box(doc, "关卡卡住了怎么办", "按 R 重置当前关卡。如果反复卡关，请检查关卡文件是否有格式错误。")
add_note_box(doc, "过关后没反应", "过关后会显示\"LEVEL COMPLETE!\"面板，按 Enter 或 Space 进入下一关。如果是最后一关，会提示\"ALL LEVELS COMPLETED!\"。")
add_note_box(doc, "如何添加自制关卡", "在项目的 levels/ 目录下创建 txt 文件，使用标准字符标记。文件名建议使用数字后缀如 level_08.txt，游戏会按数字顺序自动排序。")
add_note_box(doc, "屏幕尺寸不合适", "游戏窗口固定为 960x720 像素。可以通过桌面显示器设置的分辨率或缩放比例来获得最佳观感。")

doc.add_paragraph()

# Credits
p_credit = doc.add_paragraph()
p_credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_cr = p_credit.add_run("-" * 40)
run_cr.font.color.rgb = SECONDARY_COLOR
run_cr.font.size = Pt(10)

add_body(doc, "感谢使用推箱子游戏！")
add_body(doc, "游戏项目地址：https://github.com/qianluxi/push_box")
add_body(doc, "本手册基于 V1.1 版本编写。")
add_body(doc, "推箱子（Sokoban）由 Johan Stevanne 于 1981 年创作，是一种经典智力游戏。本项目为现代 Python 实现，仅供学习与探索。")

p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run("--- 手册结束 ---")
run_end.font.size = Pt(10)
run_end.font.name = 'Microsoft YaHei'
run_end.font.color.rgb = SECONDARY_COLOR

# Save
output_file = "push_box_game_manual.docx"
doc.save(output_file)
print(f"Generated: {output_file}")
